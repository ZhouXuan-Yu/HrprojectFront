"""Resume pipeline: text extraction (code libraries, no OCR) + DeepSeek parsing + DB ingest.

Pipeline:
    1. extract_resume_text(file_bytes, filename)  — pdfplumber / python-docx / plain text
    2. parse_resume_content(text)                 — DeepSeek structured parse, regex fallback
    3. ingest_resume(...)                         — dedup + Candidate / Resume / File records

Used by:
    - email_sync_service (IMAP 拉取简历附件自动入库)
    - /api/talent/upload-resume (人才库手动上传)
"""
import hashlib
import io
import logging
import os
import re
from datetime import datetime

log = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = ('.pdf', '.docx', '.txt', '.md', '.html', '.htm')

# Upload storage directory (backend/uploads/resumes)
_UPLOAD_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))),
    'uploads', 'resumes',
)

_BIG_COMPANIES = [
    "阿里巴巴", "腾讯", "字节跳动", "美团", "百度", "京东", "网易",
    "华为", "小米", "滴滴", "快手", "拼多多", "哔哩哔哩", "小红书",
    "微软", "谷歌", "亚马逊", "苹果", "Meta", "蚂蚁", "蚂蚁集团",
    "蔚来", "理想汽车", "小鹏", "大疆", "商汤", "科大讯飞",
]


# ===========================================================================
# Step 1 — text extraction (code libraries, replacing OCR recognition)
# ===========================================================================

def extract_resume_text(file_bytes, filename=''):
    """Extract raw text from a resume file using code libraries.

    Supports PDF (pdfplumber, PyPDF2 fallback), DOCX (python-docx, including
    table cells), and plain text / HTML files. Legacy .doc is rejected with a
    clear error (antiword/textract not bundled).

    Returns extracted text (may be empty string when nothing was readable).
    """
    ext = os.path.splitext(filename or '')[1].lower()

    if ext == '.pdf' or (not ext and file_bytes[:5] == b'%PDF-'):
        return _extract_pdf(file_bytes)
    if ext == '.docx' or file_bytes[:2] == b'PK':
        return _extract_docx(file_bytes)
    if ext in ('.txt', '.md'):
        return _decode_text(file_bytes)
    if ext in ('.html', '.htm'):
        return _strip_html(_decode_text(file_bytes))
    if ext == '.doc':
        raise ValueError('暂不支持旧版 .doc 格式，请转换为 .docx 或 .pdf 后重试')

    # Unknown extension — best-effort decode as text
    return _decode_text(file_bytes)


def _extract_pdf(file_bytes):
    """PDF text extraction: pdfplumber first (better CJK layout), PyPDF2 fallback."""
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or '')
                # Resume PDFs often put contact info in tables
                for table in page.extract_tables() or []:
                    for row in table:
                        parts.append(' '.join(cell or '' for cell in row))
        text = '\n'.join(p for p in parts if p and p.strip())
        if text.strip():
            return text
    except Exception as exc:
        log.warning("pdfplumber extraction failed: %s", exc)

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text = '\n'.join(page.extract_text() or '' for page in reader.pages)
        if text.strip():
            return text
    except Exception as exc:
        log.warning("PyPDF2 extraction failed: %s", exc)

    return ''


def _extract_docx(file_bytes):
    """DOCX text extraction via python-docx, paragraphs + table cells."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(' '.join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return '\n'.join(parts)
    except Exception as exc:
        log.warning("python-docx extraction failed: %s", exc)
        return ''


def _decode_text(file_bytes):
    """Decode bytes as text, trying common Chinese encodings."""
    for enc in ('utf-8', 'gb18030', 'gbk', 'big5', 'latin-1'):
        try:
            return file_bytes.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode('utf-8', errors='ignore')


def _strip_html(html_text):
    """Rough HTML tag strip for html-format resumes."""
    text = re.sub(r'<(script|style)[^>]*>.*?</\1>', ' ', html_text, flags=re.S | re.I)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.I)
    text = re.sub(r'</(p|div|tr|li|h\d)>', '\n', text, flags=re.I)
    text = re.sub(r'</t[dh]>', ' ', text, flags=re.I)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'&nbsp;', ' ', text)
    text = re.sub(r'&amp;', '&', text)
    return re.sub(r'\n{3,}', '\n\n', text)


# ===========================================================================
# Step 2 — DeepSeek structured parsing (with regex fallback)
# ===========================================================================

_RESUME_PARSE_SYSTEM = """你是专业的简历解析引擎。从用户提供的简历原文中提取结构化信息。

严格返回 JSON 对象，字段如下：
{
  "name": "候选人姓名（2-4个中文字符，提取不到则为空字符串）",
  "phone": "11位手机号，没有则为空字符串",
  "email": "邮箱地址，没有则为空字符串",
  "edu_level": "最高学历，只能是：大专/本科/硕士/博士 之一",
  "school_level": "院校层次，只能是：普通/211/985/C9 之一（海外名校记为C9）",
  "work_years": 数字，工作年限（应届毕业生为0，根据工作经历时间推算）,
  "recent_company": "最近一家公司名称，没有则为空字符串",
  "skills": ["核心技能列表，最多10个，如 Java、Spring Boot、MySQL"],
  "certs": ["证书列表，如 PMP、软考高级，没有则为空数组"],
  "summary": "100字以内的候选人画像摘要"
}
只返回 JSON，不要任何额外说明文字。"""


def parse_resume_content(text):
    """Parse resume text into structured data via DeepSeek.

    Falls back to the local regex-based ai_engine parser when DeepSeek is
    unavailable or returns invalid JSON. Never raises for parse-quality
    reasons — worst case returns the regex result.

    Returns dict: name, phone, email, edu_level, school_level, work_years,
                  recent_company, skills, certs, big_company_flag, summary,
                  parse_engine ('deepseek' | 'regex')
    """
    text = (text or '').strip()
    if not text:
        raise ValueError('简历内容为空，无法解析')

    # Truncate overly long resumes to control token cost
    if len(text) > 12000:
        text = text[:12000]

    try:
        from app.services import deepseek_client
        data = deepseek_client.chat_completion_json(
            messages=[
                {"role": "system", "content": _RESUME_PARSE_SYSTEM},
                {"role": "user", "content": f"简历原文：\n{text}"},
            ],
            temperature=0.1,
            max_tokens=1500,
        )
        result = _normalize_deepseek_result(data, text)
        result['parse_engine'] = 'deepseek'
        log.info("Resume parsed by DeepSeek: name=%s skills=%d",
                 result.get('name'), len(result.get('skills', [])))
        return result
    except Exception as exc:
        log.warning("DeepSeek resume parse failed, falling back to regex: %s", exc)

    from app.services.ai_engine import parse_resume as regex_parse
    result = regex_parse(text=text)
    result.setdefault('certs', [])
    result['big_company_flag'] = _detect_big_company(result.get('recent_company', ''), text)
    result['parse_engine'] = 'regex'
    return result


def _normalize_deepseek_result(data, raw_text):
    """Normalize DeepSeek JSON output to the canonical parse-result shape."""
    if not isinstance(data, dict):
        raise ValueError('DeepSeek 返回结果不是 JSON 对象')

    edu = str(data.get('edu_level') or '本科')
    if edu not in ('大专', '本科', '硕士', '博士'):
        edu = '本科'

    school = str(data.get('school_level') or '普通')
    if school not in ('普通', '211', '985', 'C9'):
        school = '普通'

    try:
        work_years = max(0, min(int(data.get('work_years') or 0), 40))
    except (TypeError, ValueError):
        work_years = 0

    skills = data.get('skills')
    if not isinstance(skills, list):
        skills = []
    skills = [str(s).strip() for s in skills if str(s).strip()][:10]

    certs = data.get('certs')
    if not isinstance(certs, list):
        certs = []
    certs = [str(c).strip() for c in certs if str(c).strip()][:10]

    name = str(data.get('name') or '').strip()
    # Sanity check: Chinese name 2-4 chars, no symbols
    if not (2 <= len(name) <= 4) or re.search(r'[a-zA-Z0-9@：:，,]', name):
        # Fall back to regex name extraction
        from app.services.ai_engine import _extract_name
        name = _extract_name(raw_text) or '未知'

    phone = str(data.get('phone') or '').strip()
    if not re.fullmatch(r'1[3-9]\d{9}', phone):
        m = re.search(r'1[3-9]\d{9}', raw_text)
        phone = m.group(0) if m else ''

    email = str(data.get('email') or '').strip()
    if not re.fullmatch(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', email):
        m = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', raw_text)
        email = m.group(0) if m else ''

    company = str(data.get('recent_company') or '').strip()

    return {
        'name': name,
        'phone': phone,
        'email': email,
        'edu_level': edu,
        'school_level': school,
        'work_years': work_years,
        'recent_company': company,
        'skills': skills,
        'certs': certs,
        'big_company_flag': _detect_big_company(company, raw_text),
        'summary': str(data.get('summary') or '').strip()[:500],
    }


def _detect_big_company(company, raw_text=None):
    """1 when the *recent company* field mentions a known big company.

    只依据 recent_company 字段匹配，不再扫描简历全文（避免把"期望入职大厂"
    之类的表述误判为大厂经历）。raw_text 参数保留仅为兼容旧调用方签名。
    """
    haystack = company or ''
    return 1 if any(c in haystack for c in _BIG_COMPANIES) else 0


# ===========================================================================
# Name fallback helpers (DeepSeek / regex 解析失败时)
# ===========================================================================

_NAME_RE = r'[一-鿿]{2,4}'


def _name_from_filename(filename):
    """从附件文件名提取姓名：张三-Java-3年.pdf / 李四_前端_5年.docx / 王五的简历.pdf"""
    stem = os.path.splitext(os.path.basename(filename or ''))[0]
    for pat in (
        rf'^({_NAME_RE})[-_–—\s]',            # 张三-Java-3年
        rf'^({_NAME_RE})的?简历',              # 张三简历 / 张三的简历
        rf'简历[-_–—\s]*({_NAME_RE})',        # 简历-张三
    ):
        m = re.search(pat, stem)
        if m:
            return m.group(1)
    return ''


def _name_from_subject(subject):
    """从邮件主题提取姓名：应聘：Java开发工程师-张三 / 张三应聘前端"""
    s = subject or ''
    for pat in (
        rf'[-_–—：:\s]({_NAME_RE})\s*$',      # 结尾 "-张三"
        rf'^({_NAME_RE})\s*(应聘|求职|申请)',  # 张三应聘…
    ):
        m = re.search(pat, s.strip())
        if m:
            return m.group(1)
    return ''


# ===========================================================================
# Target-position extraction from mail subject
# ===========================================================================

_POSITION_PATTERNS = [
    r'(?:应聘|求职|申请)\s*[:：]?\s*(?:岗位|职位)?\s*[:：]?\s*(.+?)\s*(?:岗位|职位|工程师|经理|专员|主管|总监|简历|$)',
    r'(?:岗位|职位)\s*[:：]\s*(.+?)\s*(?:简历|求职|应聘|$)',
]


def extract_target_position(subject):
    """从邮件主题解析应聘岗位名，提取不到返回 ''（不依赖 LLM）。"""
    s = (subject or '').strip()
    if not s:
        return ''
    for pat in _POSITION_PATTERNS:
        m = re.search(pat, s)
        if m:
            pos = m.group(1).strip(' -_–—：:')
            # 附带岗位后缀时补回（"Java开发" → "Java开发工程师" 语境更准）
            suffix = re.search(r'(工程师|经理|专员|主管|总监)\s*$', m.group(0))
            if suffix and suffix.group(1) not in pos:
                pos = pos + suffix.group(1)
            # 去掉结尾的候选人姓名分隔（"Java开发工程师-张三"）
            pos = re.split(r'[-_–—]', pos)[0].strip()
            if 2 <= len(pos) <= 30:
                return pos
    return ''


# ===========================================================================
# Step 3 — ingest into DB (dedup + Candidate / Resume / File records)
# ===========================================================================

_EDU_TO_INT = {'大专': 1, '本科': 2, '硕士': 3, '博士': 4}
_SCHOOL_TO_INT = {'普通': 1, '211': 2, '985': 3, 'C9': 4}


def is_resume_attachment(filename):
    """Return True when the filename looks like a resume file we can parse."""
    ext = os.path.splitext(filename or '')[1].lower()
    return ext in SUPPORTED_EXTENSIONS or ext == '.doc'


def looks_like_resume(subject, body):
    """Heuristic: does this email subject/body look like a job application?"""
    text = f'{subject or ""} {body or ""}'
    if re.search(r'简历|求职|应聘|resume|cv\b', text, re.I):
        return True
    # Body containing contact info + work experience markers
    if re.search(r'1[3-9]\d{9}', text) and re.search(r'工作经验|教育经历|项目经验', text):
        return True
    return False


def ingest_resume(file_bytes, filename, source_channel='邮箱', mail_account_id=None,
                  raw_text=None, mail_subject=None, target_position=None,
                  target_demand_id=None):
    """Full ingest pipeline: extract → parse → dedup → persist.

    Args:
        file_bytes: raw file content.
        filename: original filename (used for extension detection + storage).
        source_channel: 候选人来源（邮箱 / 手动上传 / Boss ...）.
        mail_account_id: 采集邮箱 ID（邮件渠道时传入）.
        raw_text: pre-extracted text (skips extraction when given).
        mail_subject: 邮件主题（邮件渠道时传入，写入 extract_json）.
        target_position: 从主题解析出的应聘岗位（写入 extract_json）.
        target_demand_id: 自动关联到的需求 ID（写入 extract_json）.

    Returns dict: {
        candidate_no, candidate_name, is_new_candidate, parse_engine,
        resume_id, skills, summary, target_position
    }
    """
    from app.extensions import db
    from app.models.candidate import Candidate, Resume
    from app.models.infrastructure import File

    text = raw_text if raw_text is not None else extract_resume_text(file_bytes, filename)
    if not text or not text.strip():
        raise ValueError(f'无法从文件中提取文本: {filename}')

    parsed = parse_resume_content(text)

    # ---- name fallback: 解析不出姓名时从文件名 / 邮件主题兜底 ----
    name = (parsed.get('name') or '').strip()
    if (not name or name == '未知'
            or not (2 <= len(name) <= 4)
            or re.search(r'简历|求职|应聘|个人', name)):
        name = (_name_from_filename(filename)
                or _name_from_subject(mail_subject)
                or '未知')
        parsed['name'] = name

    # ---- dedup: mobile_hash → email → 内容指纹 ----
    mobile_hash = hashlib.sha256(parsed['phone'].encode()).hexdigest() if parsed['phone'] else None
    email_hash = hashlib.sha256(parsed['email'].lower().encode()).hexdigest() if parsed['email'] else None
    # 内容指纹：正文前 2000 字符 sha256，用于无手机号/邮箱时的兜底去重
    content_fp = hashlib.sha256(text.strip()[:2000].encode('utf-8')).hexdigest()

    candidate = None
    if mobile_hash:
        candidate = Candidate.active().filter_by(mobile_hash=mobile_hash).first()
    if candidate is None and email_hash:
        # email column stores plaintext in current schema; match directly（统一 lower 兼容）
        candidate = Candidate.active().filter_by(email=parsed['email'].lower()).first()
        if candidate is None:
            candidate = Candidate.active().filter_by(email=parsed['email']).first()
    if candidate is None:
        # 内容指纹命中同候选人历史简历 → 视为重复（更新而非新建）。
        # 最近 N 份简历内做 Python 侧比对（SQLite/MySQL 通用，且该路径本就是
        # 无手机号/邮箱时的兜底，量小可接受）。
        try:
            recent = (Resume.active()
                      .order_by(Resume.id.desc()).limit(500).all())
            hit = next((r for r in recent
                        if (r.extract_json or {}).get('content_fingerprint') == content_fp),
                       None)
            if hit:
                candidate = Candidate.active().filter_by(id=hit.candidate_id).first()
        except Exception as exc:
            log.warning("content-fingerprint dedup failed (best-effort): %s", exc)

    is_new = candidate is None
    if is_new:
        candidate_no = _next_candidate_no()
        candidate = Candidate(
            candidate_no=candidate_no,
            candidate_name=parsed['name'],
            mobile=parsed['phone'] or None,
            mobile_hash=mobile_hash,
            email=parsed['email'] or None,
            source_channel=source_channel,
            status='available',
        )
        db.session.add(candidate)
    else:
        # Refresh profile fields with the latest parse
        candidate.candidate_name = parsed['name'] or candidate.candidate_name
        if parsed['phone'] and not candidate.mobile:
            candidate.mobile = parsed['phone']
            candidate.mobile_hash = mobile_hash
        if parsed['email'] and not candidate.email:
            candidate.email = parsed['email']

    # Profile fields from parse
    candidate.edu_level = _EDU_TO_INT.get(parsed['edu_level'], 2)
    candidate.school_level = _SCHOOL_TO_INT.get(parsed['school_level'], 1)
    candidate.work_years = parsed['work_years']
    candidate.big_company_flag = parsed['big_company_flag']
    candidate.cert_count = len(parsed.get('certs', []))

    from app.utils.scoring import calc_profile_score
    candidate.static_ability_score = calc_profile_score(
        edu_level=candidate.edu_level,
        school_level=candidate.school_level,
        work_years=candidate.work_years or 0,
        big_company=candidate.big_company_flag,
        cert_count=candidate.cert_count,
        skill_count=len(parsed.get('skills', [])),
    )

    db.session.flush()  # assign candidate.id

    # ---- store file ----
    file_id = None
    try:
        file_id = _store_file(file_bytes, filename, candidate.candidate_no)
    except Exception as exc:
        log.warning("Failed to store resume file %s: %s", filename, exc)

    resume = Resume(
        candidate_id=candidate.id,
        resume_file_id=file_id,
        storage_time=datetime.now(),
        base_score=candidate.static_ability_score,
        work_exp_text=parsed['summary'],
        extract_json={
            **parsed,
            'source_file': filename,
            'parsed_at': datetime.now().isoformat(timespec='seconds'),
            'content_fingerprint': content_fp,
            'mail_subject': (mail_subject or '')[:200],
            'target_position': target_position or '',
            'target_demand_id': target_demand_id,
        },
        mail_account_id=mail_account_id,
    )
    db.session.add(resume)
    db.session.commit()

    log.info(
        "Resume ingested: candidate=%s(%s) new=%s engine=%s file=%s target=%s",
        candidate.candidate_name, candidate.candidate_no, is_new,
        parsed['parse_engine'], filename, target_position or '-',
    )

    return {
        'candidate_no': candidate.candidate_no,
        'candidate_name': candidate.candidate_name,
        'is_new_candidate': is_new,
        'parse_engine': parsed['parse_engine'],
        'resume_id': resume.id,
        'skills': parsed['skills'],
        'summary': parsed['summary'],
        'target_position': target_position or '',
        'target_demand_id': target_demand_id,
    }


def _next_candidate_no():
    """Generate next candidate number: C{YYYYMM}{seq:04d}."""
    from app.models.candidate import Candidate
    prefix = 'C' + datetime.now().strftime('%Y%m')
    last = (
        Candidate.query.filter(Candidate.candidate_no.like(f'{prefix}%'))
        .order_by(Candidate.candidate_no.desc())
        .first()
    )
    seq = 1
    if last and last.candidate_no and last.candidate_no[len(prefix):].isdigit():
        seq = int(last.candidate_no[len(prefix):]) + 1
    return f'{prefix}{seq:04d}'


def _store_file(file_bytes, filename, candidate_no):
    """Persist file to uploads dir and create a File record. Returns files.id."""
    from app.extensions import db
    from app.models.infrastructure import File

    os.makedirs(_UPLOAD_DIR, exist_ok=True)
    safe_name = re.sub(r'[^\w.\-一-鿿]', '_', os.path.basename(filename or 'resume'))
    stored_name = f'{candidate_no}_{datetime.now().strftime("%H%M%S")}_{safe_name}'
    path = os.path.join(_UPLOAD_DIR, stored_name)
    with open(path, 'wb') as f:
        f.write(file_bytes)

    rec = File(
        file_name=safe_name,
        file_url=path,
        file_extension=os.path.splitext(safe_name)[1].lstrip('.').lower(),
        file_size=len(file_bytes),
        biz_type='resume',
    )
    db.session.add(rec)
    db.session.flush()
    return rec.id
